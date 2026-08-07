"""Resume → .docx (python-docx).

Why a third format at all: some application portals demand an editable Word file, and
ATS parsers handle .docx more reliably than a PDF. This output trades layout fidelity
for editability — it is deliberately the plainest of the three.

Layout fidelity is limited by what the format offers. Two consequences worth knowing:

- The "title left, date right" row is a two-cell borderless table, because Word has no
  flexbox. Right-aligning the second cell is the only reliable way to get both on one
  baseline.
- Section rules are a paragraph bottom border, which needs raw OOXML — python-docx has
  no border API. `_add_bottom_border` is that escape hatch.

Per the plan, this output gets text-layer diffing only; rendering a .docx on Linux to
compare page images would mean installing LibreOffice (+400 MB) for marginal value.
"""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from docx.text.paragraph import Paragraph

from pipeline.config import Settings
from pipeline.models import Resume
from pipeline.render.labels import (
    fmt_degree,
    fmt_education_date,
    fmt_month,
    fmt_range,
    fmt_skill,
    labels_for,
)

log = logging.getLogger(__name__)

# --- Design tokens, mirroring print.css / latex.py ---------------------------
MARGIN_X = Mm(14)
MARGIN_Y = Mm(12)

SIZE_NAME = Pt(17)
SIZE_CONTACT = Pt(9)
SIZE_SECTION = Pt(10.5)
SIZE_ITEM_TITLE = Pt(10.5)
SIZE_ITEM_META = Pt(9.5)
SIZE_BODY = Pt(9.5)

SPACE_SECTION = Pt(12)
SPACE_ITEM = Pt(7)
SPACE_BULLET = Pt(2)

INK = RGBColor(0x11, 0x11, 0x11)
INK_MUTED = RGBColor(0x44, 0x44, 0x44)

FONT_SERIF = "Noto Serif"
FONT_SANS = "Noto Sans"
FONT_CJK_SERIF = "Noto Serif CJK TC"
FONT_CJK_SANS = "Noto Sans CJK TC"

BULLET = "▪"


def render_docx(resume: Resume, settings: Settings, output: Path) -> Path:
    """Write the resume as a .docx. Returns the written path."""
    is_zh = resume.lang == "zh"
    labels = labels_for(resume.lang)
    body_font = FONT_CJK_SANS if is_zh else FONT_SANS
    heading_font = FONT_CJK_SERIF if is_zh else FONT_SERIF

    document = Document()
    _setup_page(document, body_font)

    # --- Header ---
    name = _paragraph(document, WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
    _run(name, resume.profile.name, heading_font, SIZE_NAME, bold=True, cjk=is_zh)

    if contacts := resume.profile.contact_line():
        para = _paragraph(document, WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
        _run(para, "  ·  ".join(contacts), body_font, SIZE_CONTACT, color=INK_MUTED, cjk=is_zh)

    if resume.profile.summary:
        para = _paragraph(document, WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
        _run(para, resume.profile.summary, body_font, SIZE_BODY, cjk=is_zh)

    # --- Education ---
    if education := resume.sorted_education():
        _section_heading(document, labels.education, heading_font, is_zh)
        for index, item in enumerate(education):
            _entry_row(
                document,
                item.institution,
                fmt_education_date(item, resume.lang),
                body_font,
                is_zh,
                bold_left=True,
            )
            if degree_line := fmt_degree(item, resume.lang):
                _entry_row(
                    document, degree_line, "", body_font, is_zh, italic_left=not is_zh, muted=True
                )
            _bullets(document, item.bullets, body_font, is_zh)
            _spacer(document, index, len(education))

    # --- Experience ---
    if experiences := resume.sorted_experiences():
        _section_heading(document, labels.experience, heading_font, is_zh)
        for index, item in enumerate(experiences):
            _entry_row(
                document, item.organization, item.location or "", body_font, is_zh, bold_left=True
            )
            _entry_row(
                document,
                item.role,
                fmt_range(item.start, item.end, resume.lang),
                body_font,
                is_zh,
                italic_left=not is_zh,
                muted=True,
            )
            _bullets(document, item.bullets, body_font, is_zh)
            _spacer(document, index, len(experiences))

    # --- Projects ---
    if projects := resume.sorted_projects():
        _section_heading(document, labels.projects, heading_font, is_zh)
        for index, item in enumerate(projects):
            _entry_row(
                document,
                item.name,
                item.context or fmt_month(item.date_, resume.lang),
                body_font,
                is_zh,
                bold_left=True,
            )
            _bullets(document, item.bullets, body_font, is_zh)
            _spacer(document, index, len(projects))

    # --- Publications ---
    if publications := resume.sorted_publications():
        _section_heading(document, labels.publications, heading_font, is_zh)
        for index, item in enumerate(publications):
            _entry_row(
                document,
                item.title,
                item.venue or fmt_month(item.date_, resume.lang),
                body_font,
                is_zh,
                bold_left=True,
            )
            if item.authors:
                _entry_row(
                    document,
                    item.authors,
                    "",
                    body_font,
                    is_zh,
                    italic_left=not is_zh,
                    muted=True,
                )
            _bullets(document, item.bullets, body_font, is_zh)
            _spacer(document, index, len(publications))

    # --- Skills: one paragraph per category, bold label then values ---
    if skill_groups := resume.skills_by_category():
        _section_heading(document, labels.skills, heading_font, is_zh)
        for category, skills in skill_groups.items():
            para = _paragraph(document, space_after=SPACE_BULLET)
            _run(
                para,
                f"{labels.skill_categories[category.value]}: ",
                body_font,
                SIZE_BODY,
                bold=True,
                cjk=is_zh,
            )
            values = labels.skill_join.join(fmt_skill(s, resume.lang) for s in skills)
            _run(para, values, body_font, SIZE_BODY, cjk=is_zh)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))
    log.info("rendered %s (%.1f KB)", output.name, output.stat().st_size / 1024)
    return output


# --- building blocks --------------------------------------------------------


def _setup_page(document: Document, body_font: str) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = section.right_margin = MARGIN_X
    section.top_margin = section.bottom_margin = MARGIN_Y

    style = document.styles["Normal"]
    style.font.name = body_font
    style.font.size = SIZE_BODY
    style.font.color.rgb = INK
    # Word tracks East Asian fonts in a separate attribute; without setting eastAsia the
    # CJK glyphs fall back to whatever the reader's default is.
    style.element.rPr.rFonts.set(qn("w:eastAsia"), body_font)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)


def _paragraph(
    document: Document,
    alignment: int | None = None,
    *,
    space_before: Pt | None = None,
    space_after: Pt | None = None,
) -> Paragraph:
    para = document.add_paragraph()
    if alignment is not None:
        para.alignment = alignment
    if space_before is not None:
        para.paragraph_format.space_before = space_before
    if space_after is not None:
        para.paragraph_format.space_after = space_after
    return para


def _run(
    para: Paragraph,
    text: str,
    font: str,
    size: Pt,
    *,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor = INK,
    cjk: bool = False,
) -> None:
    run = para.add_run(text)
    run.font.name = font
    run.font.size = size
    run.font.bold = bold
    # Synthesised obliques on CJK faces look broken, matching the CSS rule.
    run.font.italic = italic and not cjk
    run.font.color.rgb = color
    if cjk:
        run.element.rPr.rFonts.set(qn("w:eastAsia"), font)


def _section_heading(document: Document, text: str, font: str, is_zh: bool) -> None:
    para = _paragraph(document, space_before=SPACE_SECTION, space_after=Pt(3))
    # Uppercase is meaningless for CJK, so only the Latin heading is transformed.
    label = text if is_zh else text.upper()
    run = para.add_run(label)
    run.font.name = font
    run.font.size = SIZE_SECTION
    run.font.bold = True
    run.font.color.rgb = INK
    if is_zh:
        run.element.rPr.rFonts.set(qn("w:eastAsia"), font)
    _add_bottom_border(para)
    # Word ties "keep with next" to the paragraph, which is what stops a heading from
    # landing alone at the foot of a page.
    para.paragraph_format.keep_with_next = True


def _entry_row(
    document: Document,
    left: str,
    right: str,
    font: str,
    is_zh: bool,
    *,
    bold_left: bool = False,
    italic_left: bool = False,
    muted: bool = False,
) -> None:
    """Left text and right-aligned text on one baseline.

    A borderless two-column table, because Word has no flexbox and a right tab stop
    breaks as soon as the left text is long enough to reach it.
    """
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    _strip_table_borders(table)

    left_cell, right_cell = table.rows[0].cells
    for cell in (left_cell, right_cell):
        cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    _run(
        left_cell.paragraphs[0],
        left,
        font,
        SIZE_ITEM_TITLE if bold_left else SIZE_ITEM_META,
        bold=bold_left,
        italic=italic_left,
        color=INK_MUTED if muted else INK,
        cjk=is_zh,
    )
    if right:
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _run(right_para, right, font, SIZE_ITEM_META, color=INK_MUTED, cjk=is_zh)


def _bullets(document: Document, bullets: list[str], font: str, is_zh: bool) -> None:
    """Bullets as manually-marked paragraphs.

    Word's List Bullet style would work, but it carries the template's own indents and
    marker glyph, which differ from the other two renderers. Drawing the marker keeps
    all three consistent.
    """
    for index, text in enumerate(bullets):
        para = _paragraph(
            document,
            space_before=Pt(1) if index == 0 else SPACE_BULLET,
            space_after=Pt(0),
        )
        para.paragraph_format.left_indent = Mm(3.4)
        # Hanging indent so wrapped lines align with the text, not the marker.
        para.paragraph_format.first_line_indent = Mm(-3.4)
        _run(para, f"{BULLET} ", font, Pt(8), cjk=is_zh)
        _run(para, text, font, SIZE_BODY, cjk=is_zh)


def _spacer(document: Document, index: int, total: int) -> None:
    """Gap between entries, skipped after the last one."""
    if index < total - 1:
        _paragraph(document, space_after=SPACE_ITEM)


def _add_bottom_border(para: Paragraph) -> None:
    """Thin rule under a paragraph. python-docx has no border API, so this is raw OOXML."""
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    # w:sz is in eighths of a point; 4 ≈ 0.5pt, matching --rule-weight.
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "111111")
    borders.append(bottom)
    para._p.get_or_add_pPr().append(borders)


def _strip_table_borders(table: object) -> None:
    """Remove all borders from a layout table so it reads as plain text."""
    properties = table._tbl.tblPr  # type: ignore[attr-defined]
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "none")
        element.set(qn("w:sz"), "0")
        borders.append(element)
    properties.append(borders)
